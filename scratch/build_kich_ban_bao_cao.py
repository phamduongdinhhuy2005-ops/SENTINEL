from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "scratch/kich-ban-bao-cao-sentinel-v2.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
    return p


def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix):])
        for run in (r1, r2):
            run.font.name = "Calibri"
            run.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True)
        set_cell_shading(hdr[idx], "E8EEF5")
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(row[idx], value)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(6)
styles["Normal"].paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer_run = footer.add_run("SENTINEL v2 - Kịch bản bảo vệ dự án")
footer_run.font.name = "Calibri"
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor.from_string("555555")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(3)
run = title.add_run("KỊCH BẢN BÁO CÁO DỰ ÁN TRƯỚC HỘI ĐỒNG")
run.font.name = "Calibri"
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = RGBColor.from_string("0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(12)
run = subtitle.add_run("Dự án SENTINEL v2 - OWASP Security Workbench")
run.font.name = "Calibri"
run.font.size = Pt(13)
run.italic = True

add_para(
    doc,
    "Tài liệu này dùng để luyện trình bày, phân chia nội dung nói, chuẩn bị demo và chuẩn bị câu trả lời cho các câu hỏi hội đồng thường đặt ra khi bảo vệ dự án.",
)

add_heading(doc, "1. Thông tin cốt lõi cần nắm", 1)
core_rows = [
    ("Tên dự án", "SENTINEL v2 - công cụ desktop hỗ trợ rà soát bảo mật website và mã nguồn theo OWASP Top 10."),
    ("Vấn đề", "Sinh viên và nhóm phát triển thường tập trung chức năng nghiệp vụ nhưng bỏ sót kiểm tra header, cookie, CORS, dependency, secret, logging và các cấu hình bảo mật cơ bản."),
    ("Mục tiêu", "Xây dựng công cụ hỗ trợ phát hiện rủi ro ban đầu, phân loại findings theo OWASP, gợi ý hướng khắc phục và xuất báo cáo."),
    ("Đối tượng dùng", "Sinh viên, nhóm phát triển, người kiểm thử nội bộ hoặc nhóm cần checklist bảo mật nhanh trước khi bàn giao."),
    ("Công nghệ", "Electron, React, TypeScript, Vite, Zustand, Node.js; engine quét viết bằng JavaScript; kiểm thử bằng Vitest."),
    ("Phạm vi", "URL Scan, Project Scan, bảng findings, checklist, lịch sử quét, xuất báo cáo HTML/JSON, trợ lý AI bảo mật tùy chọn."),
    ("Giới hạn", "Không thay thế pentest thủ công; findings dạng heuristic cần xác minh thêm theo ngữ cảnh thực tế."),
]
add_table(doc, ["Mục", "Nội dung nói ngắn gọn"], core_rows, [2100, 7260])

add_heading(doc, "1.1. Kiến trúc hệ thống cần trình bày rõ", 2)
architecture_rows = [
    ("Lớp giao diện", "React + TypeScript hiển thị form quét URL, form chọn thư mục mã nguồn, bảng findings, checklist, lịch sử quét, khung trợ lý AI và nút xuất báo cáo."),
    ("Lớp desktop", "Electron main process chịu trách nhiệm tạo cửa sổ ứng dụng, truy cập tài nguyên máy người dùng và kết nối với preload bridge. Preload giúp renderer gọi chức năng hệ thống theo phạm vi kiểm soát thay vì truy cập trực tiếp Node.js."),
    ("State/UI flow", "Zustand lưu trạng thái quét, danh sách findings, lịch sử và checklist. Nhờ tách state khỏi component, giao diện có thể cập nhật tiến trình quét và kết quả theo thời gian thực mà không phải truyền props phức tạp."),
    ("Collector layer", "Các collector thu thập dữ liệu từ website hoặc dự án local: header, cookie, CORS, form, endpoint, config, dependency, secret, logging và thông tin mã nguồn liên quan."),
    ("Scanner engine", "Scan engine điều phối quá trình quét, gọi collector phù hợp, gom dữ liệu đầu vào và chuyển sang rule engine để đánh giá rủi ro."),
    ("Rule engine", "Rule engine là lõi phân tích. Mỗi rule có metadata gồm OWASP category, severity, confidence, mô tả, bằng chứng và khuyến nghị. Khi dữ liệu khớp điều kiện, rule tạo finding chuẩn hóa."),
    ("Finding model", "Finding được chuẩn hóa thành một định dạng chung gồm tên lỗi, mức độ, nhóm OWASP, vị trí, bằng chứng, độ tin cậy, trạng thái xử lý và hướng khắc phục."),
    ("Report/checklist", "Report engine xuất HTML/JSON; checklist chuyển findings thành danh sách việc cần xử lý theo mức độ ưu tiên để người dùng theo dõi quá trình khắc phục."),
    ("AI assistant", "AI assistant dùng knowledge base nội bộ và có thể gọi LLM nếu người dùng cấu hình API key. Vai trò chính là giải thích findings, OWASP và gợi ý khắc phục, không thay thế kết luận kiểm thử."),
]
add_table(doc, ["Thành phần", "Vai trò khi trình bày trước hội đồng"], architecture_rows, [2200, 7160])

add_heading(doc, "1.2. Điểm nổi bật nên nhấn mạnh", 2)
highlight_rows = [
    ("Kết hợp nhiều hướng kiểm thử", "Không chỉ quét website đang chạy mà còn phân tích mã nguồn, cấu hình và dependency. Cách kết hợp DAST + SAST + SCA giúp phát hiện rủi ro từ nhiều góc nhìn hơn."),
    ("Chuẩn hóa theo OWASP", "Findings được ánh xạ về các nhóm OWASP để hội đồng thấy dự án có cơ sở lý thuyết và chuẩn phân loại rõ ràng, không phải chỉ là danh sách lỗi rời rạc."),
    ("Có bằng chứng và khuyến nghị", "Mỗi finding không chỉ ghi tên lỗi mà còn có evidence, severity, confidence và recommendation. Đây là điểm giúp người dùng hiểu vì sao rủi ro tồn tại và nên sửa theo hướng nào."),
    ("Phù hợp môi trường học tập", "Ứng dụng desktop giúp sinh viên dễ chạy thử, quét localhost hoặc project local, xuất báo cáo và dùng làm checklist trước khi bàn giao bài hoặc sản phẩm."),
    ("Tách engine khỏi giao diện", "Phần quét/rule/report được tách khỏi React UI, nên về sau có thể mở rộng rule, thêm CLI hoặc tích hợp CI/CD mà không phải viết lại toàn bộ giao diện."),
    ("Có định hướng mở rộng", "Dự án có thể phát triển thêm verifier để giảm false positive, hỗ trợ nhiều ngôn ngữ hơn, cải thiện crawling endpoint và bổ sung AI chạy cục bộ."),
]
add_table(doc, ["Điểm nổi bật", "Cách diễn giải ngắn gọn"], highlight_rows, [2500, 6860])

add_heading(doc, "2. Kịch bản trình bày 10-12 phút", 1)
timeline_rows = [
    ("0:00-0:40", "Chào hội đồng", "Giới thiệu nhóm, tên đề tài, định hướng OWASP và mục tiêu báo cáo."),
    ("0:40-1:40", "Lý do chọn đề tài", "Nêu bối cảnh web app có nhiều điểm nhập liệu, thư viện bên thứ ba và cấu hình triển khai; lỗi nhỏ có thể dẫn đến rủi ro lớn."),
    ("1:40-2:40", "Mục tiêu và phạm vi", "Nhấn mạnh công cụ hỗ trợ rà soát ban đầu, không thay pentest; tập trung URL Scan, Project Scan, findings, checklist và báo cáo."),
    ("2:40-3:40", "Cơ sở lý thuyết", "Giải thích ngắn DAST, SAST, SCA và cách ánh xạ findings về OWASP Top 10."),
    ("3:40-5:20", "Kiến trúc hệ thống", "Trình bày theo 5 lớp: Electron desktop, React UI, Zustand state, scanner/collector, rule engine/report engine. Nói rõ dữ liệu đi từ URL/source code đến collector, qua rule engine rồi thành finding."),
    ("5:20-6:40", "Điểm nổi bật", "Nhấn mạnh DAST + SAST + SCA, chuẩn hóa theo OWASP, finding có evidence/confidence/recommendation, checklist xử lý và kiến trúc dễ mở rộng rule."),
    ("6:40-7:30", "Chức năng chính", "Đi qua URL Scan, Project Scan, quản lý findings, checklist, lịch sử quét, xuất HTML/JSON và trợ lý AI."),
    ("7:30-9:20", "Demo", "Chạy một URL hoặc mở kết quả mẫu, chỉ rõ finding, mức độ, OWASP category, bằng chứng và khuyến nghị xử lý."),
    ("9:20-10:20", "Thực nghiệm", "Nêu các kịch bản kiểm thử như OWASP Juice Shop, website có header yếu, project có dependency/cấu hình rủi ro; kết quả đạt các tiêu chí chính."),
    ("10:20-11:10", "Hạn chế", "Nói thẳng về false positive, giới hạn rule, khó hiểu logic nghiệp vụ và cần xác minh thủ công."),
    ("11:10-12:00", "Kết luận và hướng phát triển", "Tổng kết giá trị học tập/thực tiễn; định hướng mở rộng rule, hỗ trợ thêm ngôn ngữ, giảm false positive, CLI/CI và AI cục bộ."),
]
add_table(doc, ["Thời lượng", "Phần", "Ý cần nói"], timeline_rows, [1500, 1800, 6060])

add_heading(doc, "3. Lời thoại mẫu theo từng phần", 1)
script_sections = [
    ("Mở đầu", "Kính thưa quý thầy cô trong hội đồng, nhóm chúng em xin trình bày đề tài SENTINEL v2 - hệ thống hỗ trợ kiểm thử lỗ hổng bảo mật web tự động theo OWASP Top 10. Mục tiêu của nhóm là xây dựng một công cụ desktop giúp người dùng phát hiện sớm rủi ro bảo mật phổ biến trong website và mã nguồn dự án, từ đó có cơ sở ưu tiên kiểm tra và khắc phục."),
    ("Bối cảnh", "Trong quá trình phát triển ứng dụng web, các nhóm thường ưu tiên hoàn thiện chức năng như đăng nhập, quản lý dữ liệu, giỏ hàng hoặc trang quản trị. Tuy nhiên, các vấn đề như thiếu security headers, cookie chưa an toàn, CORS cấu hình sai, dependency cũ, secret bị commit hoặc log chứa dữ liệu nhạy cảm rất dễ bị bỏ sót. Vì vậy nhóm chọn đề tài này để kết hợp kiến thức lập trình, kiểm thử phần mềm và bảo mật ứng dụng web."),
    ("Mục tiêu", "SENTINEL không hướng tới thay thế công cụ pentest chuyên sâu. Dự án tập trung vào vai trò hỗ trợ rà soát ban đầu: quét URL khi ứng dụng đang chạy, phân tích mã nguồn và cấu hình dự án, chuẩn hóa findings theo OWASP, hiển thị bằng chứng, gợi ý khắc phục và cho phép xuất báo cáo."),
    ("Kiến trúc", "Về kiến trúc, SENTINEL được tổ chức theo nhiều lớp rõ ràng. Lớp ngoài cùng là ứng dụng desktop Electron, gồm main process, preload bridge và renderer. Renderer hiển thị giao diện React, còn preload đóng vai trò cầu nối an toàn để giao diện gọi các chức năng hệ thống cần thiết. Bên trong giao diện, Zustand quản lý trạng thái quét, findings, checklist và lịch sử. Phần engine phía dưới gồm collector, scanner, rule engine và report engine. Khi người dùng nhập URL hoặc chọn thư mục mã nguồn, collector thu thập dữ liệu như header, cookie, CORS, endpoint, config, dependency hoặc secret. Sau đó scanner chuyển dữ liệu sang rule engine để đánh giá. Nếu rule phát hiện tín hiệu rủi ro, hệ thống tạo finding chuẩn hóa gồm OWASP category, severity, confidence, evidence và recommendation rồi trả về giao diện."),
    ("Điểm nổi bật", "Điểm nổi bật của dự án nằm ở cách kết hợp nhiều hướng kiểm thử trong một công cụ desktop. URL Scan đóng vai trò DAST để kiểm tra website đang chạy từ bên ngoài. Project Scan đóng vai trò SAST để phân tích mã nguồn, cấu hình và logging. Phần kiểm tra dependency đóng vai trò SCA để phát hiện thư viện hoặc package có rủi ro. Ngoài việc phát hiện, hệ thống còn chuẩn hóa kết quả theo OWASP, hiển thị bằng chứng, gợi ý khắc phục, tạo checklist và xuất báo cáo. Nhờ vậy, người dùng không chỉ biết có rủi ro mà còn hiểu rủi ro thuộc nhóm nào, mức độ ưu tiên ra sao và nên xử lý theo hướng nào."),
    ("Kết luận", "Qua quá trình thực nghiệm, SENTINEL đáp ứng được các chức năng chính đã đặt ra, hoạt động ổn định trong các kịch bản thử nghiệm và phù hợp với mục tiêu học tập, rà soát nội bộ hoặc kiểm tra ban đầu trước khi bàn giao. Trong tương lai, nhóm sẽ tiếp tục mở rộng rule, nâng cao xác minh findings và phát triển khả năng tích hợp vào quy trình CI/CD."),
]
for heading, text in script_sections:
    add_heading(doc, heading, 2)
    add_para(doc, text)

add_heading(doc, "4. Kịch bản demo gợi ý", 1)
demo_rows = [
    ("1", "Mở ứng dụng", "Giới thiệu màn hình chính và các vùng chức năng: quét website, quét mã nguồn, findings, checklist, lịch sử, trợ lý AI."),
    ("2", "URL Scan", "Nhập URL demo hợp lệ, chọn phạm vi quét vừa phải, bắt đầu quét và theo dõi trạng thái xử lý."),
    ("3", "Đọc findings", "Chọn một finding tiêu biểu, giải thích tên lỗi, OWASP category, severity, evidence và recommendation."),
    ("4", "Project Scan", "Chọn thư mục dự án mẫu, chạy phân tích source/config/dependency, chỉ ra các nhóm như secret, dependency risk hoặc logging."),
    ("5", "Checklist", "Mở checklist để cho thấy hệ thống chuyển findings thành các việc cần xử lý theo mức độ ưu tiên."),
    ("6", "Xuất báo cáo", "Xuất HTML hoặc JSON, nói rằng đây là đầu ra phục vụ nộp kèm, lưu trữ hoặc chia sẻ cho nhóm phát triển."),
]
add_table(doc, ["Bước", "Thao tác", "Điểm cần nói"], demo_rows, [800, 2200, 6360])

add_heading(doc, "5. Câu hỏi hội đồng thường hỏi và câu trả lời gợi ý", 1)
qa_rows = [
    ("Vì sao chọn OWASP Top 10 làm chuẩn?", "OWASP Top 10 là danh sách rủi ro phổ biến và được dùng rộng rãi trong bảo mật ứng dụng web. Dùng OWASP giúp nhóm có khung phân loại rõ ràng, dễ giải thích findings và dễ mở rộng rule về sau."),
    ("Dự án khác gì so với các scanner có sẵn?", "SENTINEL tập trung vào mục tiêu học tập và rà soát ban đầu, kết hợp URL Scan, Project Scan, checklist, báo cáo và AI assistant trong một ứng dụng desktop đơn giản. Nhóm không khẳng định thay thế công cụ chuyên nghiệp mà tối ưu cho sinh viên và nhóm phát triển nhỏ."),
    ("DAST, SAST và SCA trong dự án được hiểu như thế nào?", "DAST kiểm tra ứng dụng đang chạy từ bên ngoài như header, cookie, CORS, endpoint. SAST phân tích mã nguồn, config, logging, secret. SCA kiểm tra dependency và file quản lý package để phát hiện thư viện cũ hoặc có rủi ro."),
    ("Kiến trúc tổng thể của SENTINEL gồm những phần nào?", "Có thể trình bày theo 5 lớp: Electron desktop layer, React UI layer, Zustand state layer, scanner/collector layer và rule/report layer. Dữ liệu đi từ URL hoặc source code vào collector, qua scanner và rule engine, sau đó được chuẩn hóa thành finding để hiển thị, tạo checklist hoặc xuất báo cáo."),
    ("Vì sao phải tách collector, scanner và rule engine?", "Collector chỉ chịu trách nhiệm thu thập dữ liệu, scanner điều phối quá trình quét, còn rule engine đánh giá rủi ro. Việc tách như vậy giúp hệ thống dễ mở rộng rule mới, dễ kiểm thử từng phần và tránh để giao diện phụ thuộc trực tiếp vào logic phân tích bảo mật."),
    ("Điểm kỹ thuật nổi bật nhất của kiến trúc là gì?", "Điểm nổi bật là hệ thống không để logic quét nằm lẫn trong giao diện. UI chỉ nhận trạng thái và findings, còn phần phân tích nằm trong engine riêng. Cách tổ chức này giúp dự án có thể phát triển thêm CLI, CI/CD hoặc nhiều loại scanner khác trong tương lai."),
    ("Làm sao hệ thống ánh xạ finding về OWASP?", "Mỗi rule được thiết kế kèm metadata như mã OWASP, severity, mô tả, evidence và recommendation. Khi rule phát hiện tín hiệu rủi ro, engine chuẩn hóa dữ liệu thành finding và gắn vào nhóm OWASP tương ứng."),
    ("Độ chính xác của findings được đảm bảo ra sao?", "Hệ thống dùng rule và heuristic nên có mức độ tin cậy khác nhau. Nhóm hiển thị confidence/evidence để người dùng tự xác minh. Các lỗi liên quan logic nghiệp vụ như phân quyền vẫn cần kiểm tra thủ công trước khi kết luận."),
    ("False positive xử lý thế nào?", "Nhóm xem false positive là hạn chế tất yếu của quét tự động. Cách giảm là kết hợp nhiều tín hiệu, ghi rõ bằng chứng, phân loại confidence, bổ sung cơ chế verifier và cho phép người dùng đánh dấu trạng thái xử lý."),
    ("Vì sao dùng Electron thay vì web thuần?", "Electron giúp đóng gói thành ứng dụng desktop, thuận tiện chọn thư mục mã nguồn local và chạy trên Windows. Điều này phù hợp với mục tiêu quét project trên máy người dùng mà không cần triển khai server riêng."),
    ("Vì sao dùng React, Vite và Zustand?", "React giúp xây dựng giao diện component rõ ràng, Vite hỗ trợ phát triển nhanh, Zustand nhẹ và phù hợp quản lý state findings, lịch sử, checklist và trạng thái quét mà không làm kiến trúc quá nặng."),
    ("Ứng dụng có an toàn khi quét URL nội bộ không?", "Trong môi trường học tập, ứng dụng có thể quét target nội bộ như localhost để demo. Khi dùng thực tế cần chỉ quét hệ thống mình sở hữu hoặc được phép kiểm tra. Có thể cấu hình để hạn chế private target trong production."),
    ("AI assistant có vai trò gì?", "AI là tính năng hỗ trợ giải thích findings, OWASP và gợi ý khắc phục. Nếu không có API key, app vẫn có knowledge base nội bộ. Nếu có key, app có thể gọi LLM để trả lời sâu hơn, nhưng kết quả AI vẫn cần kiểm tra lại."),
    ("Kết quả thực nghiệm chứng minh điều gì?", "Thực nghiệm cho thấy hệ thống chạy được các kịch bản chính: quét URL, quét mã nguồn, hiển thị findings, tạo checklist và xuất báo cáo. Điều này chứng minh tính khả thi của công cụ trong phạm vi rà soát ban đầu."),
    ("Hạn chế lớn nhất của dự án là gì?", "Hạn chế lớn nhất là hệ thống chưa hiểu đầy đủ logic nghiệp vụ và ngữ cảnh thực tế của từng ứng dụng, nên các nhóm như Broken Access Control, Insecure Design hoặc SSRF cần xác minh thủ công."),
    ("Nếu phát triển tiếp, nhóm ưu tiên gì?", "Ưu tiên mở rộng rule, bổ sung verifier để giảm cảnh báo sai, hỗ trợ thêm PHP/Java/Python, cải thiện crawling endpoint, thêm CLI/CI để tích hợp GitHub Actions hoặc GitLab CI, và nghiên cứu AI cục bộ."),
    ("Nếu hội đồng hỏi dự án có thể thương mại hóa không?", "Có thể phát triển theo hướng công cụ hỗ trợ developer nội bộ hoặc giáo dục bảo mật. Tuy nhiên để thương mại hóa cần tăng độ chính xác, chuẩn hóa báo cáo, thêm quản lý người dùng, policy, cập nhật CVE và kiểm thử trên nhiều hệ thống thực tế."),
    ("Nếu hỏi tại sao không dùng mô hình AI để tự phát hiện mọi lỗi?", "AI có thể hỗ trợ phân tích và giải thích, nhưng phát hiện lỗi bảo mật cần bằng chứng, tính lặp lại và kiểm soát sai số. Vì vậy nhóm dùng rule-based engine làm lõi, AI đóng vai trò hỗ trợ thay vì thay thế hoàn toàn quá trình kiểm thử."),
]
add_table(doc, ["Câu hỏi", "Câu trả lời gợi ý"], qa_rows, [3300, 6060])

add_heading(doc, "6. Câu hỏi khó cần chuẩn bị kỹ", 1)
hard_rows = [
    ("Rule của nhóm dựa trên nguồn nào?", "Trả lời rằng rule được xây dựng theo OWASP Top 10, OWASP Testing Guide, kinh nghiệm từ các lỗ hổng phổ biến và thử nghiệm trên ứng dụng mẫu. Nếu có rule tự thiết kế, nói rõ đó là heuristic hỗ trợ cảnh báo ban đầu."),
    ("Có chứng minh exploit được không?", "Nói rằng phạm vi hiện tại chủ yếu phát hiện tín hiệu rủi ro và bằng chứng cấu hình/source, chưa tự động khai thác chuyên sâu. Hướng phát triển là thêm verifier hoặc PoC có kiểm soát cho một số lỗi an toàn."),
    ("Quét có gây ảnh hưởng target không?", "Nhấn mạnh công cụ chỉ nên dùng trên hệ thống được phép. Cường độ quét cần giới hạn; không dùng payload phá hoại; các thử nghiệm nên chạy ở môi trường local hoặc staging."),
    ("Vì sao kết quả quét có thể khác công cụ khác?", "Do mỗi công cụ có rule, crawler, payload, phạm vi và cách đánh giá severity khác nhau. SENTINEL ưu tiên dễ hiểu, phù hợp học tập và checklist ban đầu."),
    ("Hội đồng yêu cầu demo lỗi cụ thể", "Chuẩn bị sẵn một finding dễ giải thích như missing security header, cookie thiếu flag, CORS yếu, dependency cũ hoặc secret/config risk. Trình bày theo công thức: dấu hiệu, vì sao nguy hiểm, bằng chứng, cách sửa."),
]
add_table(doc, ["Tình huống", "Cách trả lời"], hard_rows, [3300, 6060])

add_heading(doc, "7. Công thức trả lời nhanh khi bị hỏi bất ngờ", 1)
formula_rows = [
    ("Bước 1", "Nhắc lại câu hỏi bằng một câu ngắn để xác nhận đúng ý."),
    ("Bước 2", "Trả lời trực tiếp trước, tránh vòng vo."),
    ("Bước 3", "Gắn câu trả lời với dự án SENTINEL: module nào, rule nào, dữ liệu nào."),
    ("Bước 4", "Nêu giới hạn nếu có, vì hội đồng thường đánh giá cao sự trung thực kỹ thuật."),
    ("Bước 5", "Chốt bằng hướng cải thiện hoặc cách kiểm chứng."),
]
add_table(doc, ["Bước", "Cách làm"], formula_rows, [1200, 8160])
add_para(doc, "Mẫu câu: Về ý này, trong phạm vi dự án nhóm xử lý ở mức ..., bằng cách .... Điểm mạnh là ..., tuy nhiên hạn chế hiện tại là .... Nếu phát triển tiếp, nhóm sẽ ....")

add_heading(doc, "8. Checklist trước ngày bảo vệ", 1)
check_rows = [
    ("Slide", "Kiểm tra không quá nhiều chữ, có sơ đồ kiến trúc, luồng xử lý, ảnh demo và kết quả thực nghiệm."),
    ("Demo", "Chuẩn bị dữ liệu mẫu, URL/project mẫu, file báo cáo xuất sẵn và phương án dự phòng nếu mạng hoặc app lỗi."),
    ("Phân vai", "Mỗi thành viên nắm phần mình nói và ít nhất 3 câu hỏi kỹ thuật liên quan phần đó."),
    ("Số liệu", "Nhớ các con số chính: số kịch bản thử nghiệm, nhóm OWASP phát hiện, công nghệ sử dụng và giới hạn của hệ thống."),
    ("Tinh thần trả lời", "Trả lời ngắn, có bằng chứng, không khẳng định quá mức; câu nào chưa chắc thì nêu hướng kiểm chứng."),
]
add_table(doc, ["Hạng mục", "Cần chuẩn bị"], check_rows, [2100, 7260])

doc.add_section(WD_SECTION.NEW_PAGE)
add_heading(doc, "Phụ lục: Mẫu kết luận 30 giây", 1)
add_para(doc, "Tóm lại, SENTINEL v2 đã hoàn thành mục tiêu xây dựng một công cụ desktop hỗ trợ rà soát bảo mật website và mã nguồn theo OWASP Top 10. Hệ thống giúp người dùng phát hiện rủi ro ban đầu, phân loại findings, theo dõi checklist và xuất báo cáo. Dù còn hạn chế trong việc xác minh logic nghiệp vụ và giảm false positive, dự án cho thấy tính khả thi trong học tập, kiểm tra nội bộ và là nền tảng để phát triển tiếp thành công cụ kiểm thử bảo mật hoàn chỉnh hơn.")

doc.save(OUT)
print(OUT)
